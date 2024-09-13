import os
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from bs4 import BeautifulSoup
import requests
import pandas as pd
from selenium.webdriver.support.ui import Select
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 
from selenium.common.exceptions import ElementClickInterceptedException
lists=[]
doc = Document()
table = doc.add_table(rows=1490, cols=8)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Info 1'
hdr_cells[1].text = 'Info 2'
hdr_cells[2].text = 'Info 3'
hdr_cells[3].text = 'Info 4'
hdr_cells[4].text = 'Info 5'
hdr_cells[5].text = 'Info 6'
hdr_cells[6].text = 'Info 7'
hdr_cells[7].text = 'Info 8'

x345 = input("enter sub category")


def _login_and_visit_website(driver, url, email, password):
    driver.get(url)
    print("Visiting the website...")

    # Logging in
    email_input = driver.find_element(By.ID, "email")
    password_input = driver.find_element(By.ID, "password")
    email_input.send_keys(email)
    password_input.send_keys(password)


    submit_btn = driver.find_element(By.CSS_SELECTOR, "button[type='submit']")
    submit_btn.click()

    print("Logged in successfully!")
    time.sleep(10)
    url=driver.current_url
    driver.get(url)
    time.sleep(3)
    return driver
def searchp(driver):
    search_program_link = driver.find_element(By.LINK_TEXT, "Search Program")
    search_program_link.click()
    time.sleep(5)
    url=driver.current_url
    driver.get(url)
    time.sleep(3)
    return driver
def advancep(driver):
    #x23=input("enter Sub Category")

    #advance_input=driver.find_element(By.ID,"txtsearch")
    #advance_input.send_keys("business")


    dropdown2 = (driver.find_element(By.CSS_SELECTOR, "[data-id='IntakeList']"))
    #time.sleep(2)
    #dropdown2.click()
    #all = driver.find_element(By.LINK_TEXT,"All")
    #time.sleep(2)
    #all.click()
    #time.sleep(2)

    '''dropdown3 = (driver.find_element(By.CSS_SELECTOR, "[data-id='Year']"))
    time.sleep(2)
    dropdown3.click()
    time.sleep(2)
    year2 = driver.find_element(By.LINK_TEXT,'2024')
    time.sleep(2)
    year2.click()'''


    advance_search_btn = driver.find_element(By.ID,"aadvsrc")
    advance_search_btn.click()
    #url=driver.current_url
    #driver.get(url)
    time.sleep(5)
    #dropdown = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "Country")))
    
    dropdown = (driver.find_element(By.CSS_SELECTOR, "[data-id='Country']"))
    time.sleep(2)
    dropdown.click()
    time.sleep(2)


    #dropdown.select_by_visible_text("United States of America")
    united = driver.find_element(By.LINK_TEXT,"United States of America")
    united.click()
    time.sleep(5)

    h2_element = driver.find_element(By.TAG_NAME,"h2")
    h2_element.click()
    time.sleep(3)


    dropdown = (driver.find_element(By.CSS_SELECTOR, "[data-id='SubjectCategory']"))
    time.sleep(2)
    dropdown.click()
    time.sleep(2)
    sub = driver.find_element(By.LINK_TEXT,"Health")
    sub.click()
    time.sleep(5)

    h2_element = driver.find_element(By.TAG_NAME,"h2")
    h2_element.click()
    time.sleep(3)

    dropdown = (driver.find_element(By.CSS_SELECTOR, "[data-id='SubjectSubCategory']"))
    time.sleep(2)
    dropdown.click()
    time.sleep(2)
    
    sub = driver.find_element(By.LINK_TEXT,x345)
    sub.click()
    time.sleep(5)

    h2_element = driver.find_element(By.TAG_NAME,"h2")
    h2_element.click()
    time.sleep(5)

    searchbtn=driver.find_element(By.ID,"Searchbtn")
    searchbtn.click()
    time.sleep(5)


    return driver

def get_links(driver):
    i=0
    while True:
        try:
            i=i+1
            container=driver.find_element(By.ID,"searchdta")
            container_html = container.get_attribute('outerHTML')
            soup = BeautifulSoup(container_html, 'html.parser')
            links = soup.find_all('a', class_='u-link-v5 text-black font-size16 font-bold')
            for link in links:
                link1=link.get('href')
                full_link = "https://www.coursefinder.ai" + link1  # Add base path to the link
                print(full_link)
                lists.append(full_link)
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            btn123 = driver.find_element(By.CSS_SELECTOR, ".fa-step-forward")
            btn123.click()
            time.sleep(3)#3
            if i==5:
                break
        except ElementClickInterceptedException:
            break  # Break out of the loop if the button is not clickable
    time.sleep(5)
    return driver
def text_extract(driver,test_url,no):
    url=test_url
    driver.get(url)
    time.sleep(3)
    span_element = driver.find_element(By.CSS_SELECTOR, ".text-blue.font-size18.font-bold")
    university_name = span_element.text
    print(university_name)

    div_element = driver.find_element(By.CSS_SELECTOR,".panel-heading.font-bold.text-white")
    subject_name=div_element.text
    print(subject_name)
    container = driver.find_element(By.CSS_SELECTOR,"div.panel.panel-info1.margin-bottom-0")
    #time.sleep(2)#2
    container_html = container.get_attribute('outerHTML')
    #time.sleep(2)#3
    # Parse the HTML content of the container using BeautifulSoup
    soup = BeautifulSoup(container_html, 'html.parser')
    #time.sleep(2)#5
    english_proficiency_section = soup.find('div', {'class': 'panel-heading'}, text='English Proficiency Test Requirements')

    # Extract the list of English proficiency test requirements

    test_info = [] 
    
    
    if english_proficiency_section:
        proficiency_tests = english_proficiency_section.find_next('ul', {'class': 'list-unstyled'})
        if proficiency_tests:
            for test in proficiency_tests.find_all('li'):
                test_name_elem = test.find('div', {'class': 'font-bold'})
                test_score_elem = test.find('div', {'class': 'align-top'})
            
                if test_name_elem and test_score_elem:
                    test_name = test_name_elem.text.strip()
                    test_score = test_score_elem.text.strip()
                    print(f"{test_name}: {test_score}")
                    test_info.append(f"{test_name}: {test_score}")
                else:
                    test_name = test_name_elem.text.strip()
                    print(f"{test_name}")
                    test_info.append(f"{test_name}")
    else:
        print("No proficiency tests found in the English proficiency section.")


    container = driver.find_element(By.CSS_SELECTOR,"div.panel.panel-info1.margin-bottom-0")
    #time.sleep(2)#2
    container_html = container.get_attribute('outerHTML')
    #time.sleep(2)#3
    # Parse the HTML content of the container using BeautifulSoup
    soup = BeautifulSoup(container_html, 'html.parser')
    #time.sleep(2)#5
    standard_requirements = soup.find('div', {'class': 'panel-heading'}, text='Standardized Test Requirements')

    # Extract the list of English proficiency test requirements
    standard_info = []

    if standard_requirements:
        proficiency_tests = standard_requirements.find_next('ul', {'class': 'list-unstyled'})
        for test in proficiency_tests.find_all('li'):
            test_name = test.find('div', {'class': 'font-bold'}).text.strip()
            if(test.find('div', {'class': 'align-top'})):
                test_score = test.find('div', {'class': 'align-top'}).text.strip()
                print(f"{test_name}: {test_score}")
                standard_info.append(f"{test_name}: {test_score}")
            else:
                print(f"{test_name}")
                standard_info.append(f"{test_name}")
    else:
        print(" not found.")

    #to find application deadline
    all_li_elements = soup.find_all('li', class_='d-flex align-items-center justify-content-between g-brd-bottom g-brd-gray-light-v4 padding-top-5 padding-bottom-5')
    deadline_text_without_keyword=""
    for li_element in all_li_elements:
        if "Application Deadline" in li_element.get_text():
            deadline_text = li_element.get_text(strip=True)
            # Exclude "Application Deadline" from the text
            deadline_text_without_keyword = deadline_text.replace("Application Deadline", "").strip()
            print(deadline_text_without_keyword)
    
    application_fee = ""
    for li_element in all_li_elements:
        if "Application Fee" in li_element.get_text():
            deadline_text = li_element.get_text(strip=True)
            # Exclude "Application Deadline" from the text
            application_fee = deadline_text.replace("Application Fee", "").strip()
            print(application_fee)

    for li_element in all_li_elements:
        if "Yearly Tuition Fee" in li_element.get_text():
            deadline_text = li_element.get_text(strip=True)
            # Exclude "Application Deadline" from the text
            year_fee = deadline_text.replace("Yearly Tuition Fee", "").strip() 
            print(year_fee)
    campus = ""
    for li_element in all_li_elements:
        if "Campus" in li_element.get_text():
            deadline_text = li_element.get_text(strip=True)
            # Exclude "Application Deadline" from the text
            campus = deadline_text.replace("Campus", "").strip()
            print(campus)

    inf1=university_name
    inf2=subject_name
    inf3=test_info
    inf4=standard_info
    inf5=deadline_text_without_keyword
    inf6=application_fee
    inf7=year_fee
    inf8=campus

    cells = table.rows[no].cells
    for i, info in enumerate([inf1, inf2, inf3, inf4, inf5, inf6, inf7, inf8]):
        if i==2:
            cell=cells[i]
            for x in inf3:
                cell.add_paragraph(x)
        else:
            cell = cells[i]
            cell.text = info

        # Set visible borders for each cell
        set_cell_border(cell, top='single', bottom='single', left='single', right='single')
    return driver


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create a new element for cell borders
    tcBorders = OxmlElement('w:tcBorders')

    # Set border properties
    for key, value in kwargs.items():
        if value:
            tag = 'w:{}'.format(key)
            border_element = OxmlElement(tag)
            border_element.set(qn('w:val'), value)
            tcBorders.append(border_element)

    # Apply borders to the cell
    tcPr.append(tcBorders)

def main():
    # Ask the user to input the path to the locally hosted website served by live-server
    path = input("enter the path ")
    web_page1_url = f"{path}"

    # Configure Selenium webdriver options
    options = webdriver.ChromeOptions()
    #options.add_argument("--headless")  # Run Chrome in headless mode (without opening browser window)

    # Initialize the ChromeDriver service with ChromeDriverManager
    
    service = Service("C:/Users/vishwa/OneDrive/Desktop/agency/chromedriver.exe")

    # Initialize the webdriver with the specified service and options
    driver = webdriver.Chrome(service=service, options=options)

    # Perform login and visit the website
    email = "useru@gmail.com"
    password = "password"
    driver = _login_and_visit_website(driver, web_page1_url, email, password)
    driver = searchp(driver)
    driver =advancep(driver)
    driver=get_links(driver)
    print("these are links that are retrieved")
    print(lists)
    print(len(lists))
    no=0
    for link in lists:
        no=no+1
        driver = text_extract(driver,link,no)
    
    # Save the Word document with multiple rows of information
    doc.save('archi.docx')

    


if __name__ == "__main__":
    main()
